from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Caminho da logo
caminho_logo = r'C:\Users\victo\OneDrive\Documentos\projeto-termometro\app\static\logo.png'

# Pasta com os QR codes (nomes dos arquivos são as identificações)
caminho_qr = r'C:\Users\victo\OneDrive\Documentos\projeto-termometro\app\static\qrcodes'
identificacoes = [os.path.splitext(f)[0] for f in os.listdir(caminho_qr) if f.endswith(('.png', '.jpg', '.jpeg'))]

# Cria documento Word
documento = Document()

# Ajusta para papel Carta
section = documento.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(0.3)
section.bottom_margin = Inches(0.3)
section.left_margin = Inches(0.3)
section.right_margin = Inches(0.3)

# Configurações de etiquetas
colunas = 4
linhas = 20
etiquetas_por_pagina = colunas * linhas

# Gerar as etiquetas em blocos
for i in range(0, len(identificacoes), etiquetas_por_pagina):
    bloco = identificacoes[i:i + etiquetas_por_pagina]
    tabela = documento.add_table(rows=linhas, cols=colunas)
    tabela.autofit = False

    idx = 0
    for r in range(linhas):
        for c in range(colunas):
            cell = tabela.cell(r, c)
            cell.width = Inches(2.0)  # Aproximadamente 4 colunas em 8.5"
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if idx < len(bloco):
                # Logo
                if os.path.exists(caminho_logo):
                    run = p.add_run()
                    run.add_picture(caminho_logo, width=Inches(0.6))
                    p.add_run("\n")
                
                # Identificação
                run = p.add_run(bloco[idx])
                run.font.name = 'Arial'
                run.bold = True
                run.font.size = Pt(9)
                idx += 1

    documento.add_page_break()

# Salvar
documento.save("etiquetas_identificacao.docx")
print("Arquivo salvo como etiquetas_identificacao.docx")
