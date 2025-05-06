from docx import Document
from docx.shared import Cm
from PIL import Image
import os

# Caminho da pasta com QR codes
path_qrcodes = r"C:\Users\victo\OneDrive\Documentos\projeto-termometro\app\static\qrcodes"

# Verifica se a pasta existe
if not os.path.exists(path_qrcodes):
    print(f"❌ Pasta '{path_qrcodes}' não encontrada.")
    exit()

# Lista arquivos .png
arquivos = sorted([f for f in os.listdir(path_qrcodes) if f.endswith(".png")])

if not arquivos:
    print("❌ Nenhum arquivo PNG encontrado na pasta de QR codes.")
    exit()

# Criar documento Word
doc = Document()

for nome_arquivo in arquivos:
    try:
        id_termometro = nome_arquivo.replace(".png", "")
        img_path = os.path.join(path_qrcodes, nome_arquivo)

        # Verifica se a imagem é válida
        Image.open(img_path).verify()

        doc.add_paragraph(f"ID: {id_termometro}")
        doc.add_picture(img_path, width=Cm(5), height=Cm(5))
        doc.add_paragraph("\n")

    except Exception as e:
        print(f"❌ Erro ao processar '{nome_arquivo}': {e}")

# Salvar documento
doc.save("etiquetas_termometros.docx")
print("✅ Documento gerado com sucesso!")
