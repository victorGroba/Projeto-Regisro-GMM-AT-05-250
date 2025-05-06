from docx import Document
from docx.shared import Cm, Pt
import os

# Caminho da pasta com os QR codes
caminho_qr = r"C:\Users\victo\OneDrive\Documentos\projeto-termometro\app\static\qrcodes"

# Cria o documento Word
documento = Document()

# Lista de imagens válidas
imagens = [img for img in os.listdir(caminho_qr) if img.lower().endswith(('.png', '.jpg', '.jpeg'))]

# Tamanho da imagem (2,5 x 2,5 cm)
tamanho_img = Cm(2.5)

# Número de linhas e colunas por página
linhas = 7
colunas = 3
por_pagina = linhas * colunas

# Processa em blocos de 21 (7x3)
for i in range(0, len(imagens), por_pagina):
    bloco = imagens[i:i+por_pagina]
    tabela = documento.add_table(rows=linhas, cols=colunas)

    idx = 0
    for r in range(linhas):
        for c in range(colunas):
            cell = tabela.cell(r, c)
            cell.text = ''
            paragrafo = cell.paragraphs[0]
            paragrafo.alignment = 1  # Centraliza

            if idx < len(bloco):
                nome_arquivo = bloco[idx]
                identificacao = os.path.splitext(nome_arquivo)[0]
                caminho_img = os.path.join(caminho_qr, nome_arquivo)

                run = paragrafo.add_run()
                run.add_picture(caminho_img, width=tamanho_img, height=tamanho_img)

                # Nome abaixo
                run = paragrafo.add_run(f"\n{identificacao}")
                run.font.size = Pt(8)
                idx += 1
            else:
                paragrafo.add_run("")  # célula vazia

    documento.add_page_break()

# Salva o documento final
documento.save("qrcodes_identificados_max_por_pagina.docx")
